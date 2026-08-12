"""Bounded text extraction for Official References (no OCR).

Supports TXT/MD via stdlib; PDF via pypdf; DOCX via python-docx when installed.
Scanned/image-only PDFs return empty text (OCR deferred).
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

# Hard caps — keep extraction cheap and safe.
_MAX_CHARS = 12_000
_MAX_PDF_PAGES = 3
_MAX_DOCX_PARAS = 80
_MAX_READ_BYTES = 256 * 1024

EXTRACT_SUFFIXES = frozenset({".txt", ".md", ".pdf", ".docx"})


def extract_text_from_path(path: Path | str, *, filename: str = "") -> dict[str, Any]:
    """Return ``{text, ok, reason, suffix}`` without raising for unsupported files."""
    target = Path(path)
    name = filename or target.name
    suffix = Path(name).suffix.lower() or target.suffix.lower()
    if suffix not in EXTRACT_SUFFIXES:
        return {
            "text": "",
            "ok": False,
            "reason": "unsupported_type",
            "suffix": suffix,
        }
    if not target.is_file():
        return {"text": "", "ok": False, "reason": "missing_file", "suffix": suffix}

    try:
        if suffix in {".txt", ".md"}:
            text = _extract_plain(target)
        elif suffix == ".pdf":
            text = _extract_pdf(target)
        elif suffix == ".docx":
            text = _extract_docx(target)
        else:
            text = ""
    except Exception:  # noqa: BLE001 — never break upload on extract failure
        return {"text": "", "ok": False, "reason": "extract_failed", "suffix": suffix}

    text = _normalize_whitespace(text)[:_MAX_CHARS]
    if not text.strip():
        return {
            "text": "",
            "ok": False,
            "reason": "empty_or_scanned",
            "suffix": suffix,
        }
    return {"text": text, "ok": True, "reason": "", "suffix": suffix}


def extract_text_from_bytes(
    data: bytes,
    *,
    filename: str,
) -> dict[str, Any]:
    """Extract from an in-memory upload (writes a temp-less path via suffix dispatch)."""
    suffix = Path(filename or "").suffix.lower()
    if suffix not in EXTRACT_SUFFIXES:
        return {
            "text": "",
            "ok": False,
            "reason": "unsupported_type",
            "suffix": suffix,
        }
    if not data:
        return {"text": "", "ok": False, "reason": "empty_or_scanned", "suffix": suffix}
    # Bound inbound bytes before decode/parse.
    blob = data[:_MAX_READ_BYTES]
    try:
        if suffix in {".txt", ".md"}:
            text = blob.decode("utf-8", errors="replace")
        elif suffix == ".pdf":
            text = _extract_pdf_bytes(blob)
        elif suffix == ".docx":
            text = _extract_docx_bytes(blob)
        else:
            text = ""
    except Exception:  # noqa: BLE001
        return {"text": "", "ok": False, "reason": "extract_failed", "suffix": suffix}

    text = _normalize_whitespace(text)[:_MAX_CHARS]
    if not text.strip():
        return {
            "text": "",
            "ok": False,
            "reason": "empty_or_scanned",
            "suffix": suffix,
        }
    return {"text": text, "ok": True, "reason": "", "suffix": suffix}


def _normalize_whitespace(text: str) -> str:
    # Keep newlines for subject-line scanning; collapse extreme runs.
    lines = [re_sub_spaces(line) for line in str(text or "").replace("\r\n", "\n").split("\n")]
    return "\n".join(lines).strip()


def re_sub_spaces(line: str) -> str:
    import re

    return re.sub(r"[ \t]+", " ", line).strip()


def _extract_plain(path: Path) -> str:
    raw = path.read_bytes()[:_MAX_READ_BYTES]
    return raw.decode("utf-8", errors="replace")


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        return ""
    try:
        reader = PdfReader(str(path), strict=False)
        parts: list[str] = []
        for page in reader.pages[:_MAX_PDF_PAGES]:
            try:
                parts.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001
                continue
            if sum(len(p) for p in parts) >= _MAX_CHARS:
                break
        return "\n".join(parts)
    except Exception:  # noqa: BLE001 — corrupt / image-only / truncated PDFs
        return ""


def _extract_pdf_bytes(data: bytes) -> str:
    try:
        from io import BytesIO

        from pypdf import PdfReader  # type: ignore
    except ImportError:
        return ""
    try:
        reader = PdfReader(BytesIO(data), strict=False)
        parts: list[str] = []
        for page in reader.pages[:_MAX_PDF_PAGES]:
            try:
                parts.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001
                continue
            if sum(len(p) for p in parts) >= _MAX_CHARS:
                break
        return "\n".join(parts)
    except Exception:  # noqa: BLE001
        return ""



def _extract_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return ""
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs[:_MAX_DOCX_PARAS]:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
        if sum(len(p) for p in parts) >= _MAX_CHARS:
            break
    return "\n".join(parts)


def _extract_docx_bytes(data: bytes) -> str:
    try:
        from io import BytesIO

        from docx import Document  # type: ignore
    except ImportError:
        return ""
    doc = Document(BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs[:_MAX_DOCX_PARAS]:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
        if sum(len(p) for p in parts) >= _MAX_CHARS:
            break
    return "\n".join(parts)
